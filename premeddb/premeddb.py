from flask import Flask
from flask_bootstrap import Bootstrap
from flask_sqlalchemy import SQLAlchemy
from flask import render_template, request, redirect, url_for, send_file # imports rendering functions
from docx import Document
from io import BytesIO
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, BooleanField
from wtforms.validators import InputRequired, Email, Length
#from models import *
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_admin import Admin
from flask_admin.contrib.sqla import ModelView

# Flask app defined
app = Flask(__name__)

# Establishes secret key
app.config['SECRET_KEY'] = 'thisisasecret'
Bootstrap(app)

# Links to database which is created in config.py
app.config.from_pyfile('config.py')
db = SQLAlchemy(app)
admin = Admin(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'


#models
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(15), unique=True)
    email = db.Column(db.String(50), unique=True)
    password = db.Column(db.String(80))
    admin = db.Column(db.String(1))
    scheduler = db.relationship('Scheduler', backref='student', lazy='dynamic')
    mcat = db.relationship('Mcat', backref='student', lazy='dynamic')
    grades = db.relationship('Grades', backref='student', lazy = 'dynamic')
    references = db.relationship('References', backref='student', lazy='dynamic')
    activities = db.relationship('Activities', backref='student', lazy='dynamic')
    status = db.relationship('Status', backref='student', lazy='dynamic')
    personal = db.relationship('Personal', backref='student', lazy='dynamic')

    def __repr__(self):
        return (self.username)


class Scheduler(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    schedulename = db.Column(db.String(50))
    schedule = db.Column(db.String(200))
    data = db.Column(db.LargeBinary)


class Grades(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    ogpa = db.Column(db.String(5))
    sgpa = db.Column(db.String(5))

    def __repr__(self):
        return 'User %r' % (self.userid)


class Mcat(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    examdate = db.Column(db.String(15))
    overall = db.Column(db.String(5))
    cp = db.Column(db.String(5))
    cars = db.Column(db.String(5))
    bb = db.Column(db.String(5))
    ps = db.Column(db.String(5))

    def __repr__(self):
        return 'User %r' % (self.userid)


class References(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    name = db.Column(db.String(50))
    email = db.Column(db.String(50))
    type = db.Column(db.String(50))
    status = db.Column(db.String(500))


class Activities(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    activity = db.Column(db.String(50))
    type = db.Column(db.String(50))
    hours = db.Column(db.String(50))
    reference = db.Column(db.String(50))
    startdate = db.Column(db.String(50))
    enddate = db.Column(db.String(50))
    description = db.Column(db.String(10000))


class Status(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    university = db.Column(db.String(50))
    primary = db.Column(db.String(50))
    secondary = db.Column(db.String(50))
    interview = db.Column(db.String(50))
    offer = db.Column(db.String(50))
    essay1p = db.Column(db.String(500))
    essay1a = db.Column(db.String(10000))
    essay2p = db.Column(db.String(500))
    essay2a = db.Column(db.String(10000))
    essay3p = db.Column(db.String(500))
    essay3a = db.Column(db.String(10000))
    essay4p = db.Column(db.String(500))
    essay4a = db.Column(db.String(10000))
    essay5p = db.Column(db.String(500))
    essay5a = db.Column(db.String(10000))


class Personal(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    title = db.Column(db.String(50))
    essay = db.Column(db.String(10000))

admin.add_view(ModelView(User, db.session))
admin.add_view(ModelView(Scheduler, db.session))
admin.add_view(ModelView(Grades, db.session))
admin.add_view(ModelView(Mcat, db.session))
admin.add_view(ModelView(References, db.session))
admin.add_view(ModelView(Activities, db.session))
admin.add_view(ModelView(Status, db.session))
admin.add_view(ModelView(Personal, db.session))


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# Creates class for login form w/ username and password
class LoginForm(FlaskForm):
    username = StringField('username', validators=[InputRequired(), Length(min=4, max=15)])
    password = PasswordField('password', validators=[InputRequired(), Length(min=8, max=80)])
    remember = BooleanField('remember me')


# Creates class for registration form w/ username, email, and password
class RegisterForm(FlaskForm):
    username = StringField('username', validators=[InputRequired(), Length(min=4, max=15)])
    email = StringField('email', validators=[InputRequired(), Email(message='Invalid email'), Length(max=50)])
    password = PasswordField('password', validators=[InputRequired(), Length(min=8, max=80)])


@app.route('/')
def index():
    return render_template("index.html")


@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()

    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user:
            if check_password_hash(user.password, form.password.data):
                login_user(user, remember=form.remember.data)
                return redirect(url_for('dashboard'))

            return "<h1>Username or password is not valid</h1>"

    return render_template("login.html", form=form)


@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    form = RegisterForm()

    if form.validate_on_submit():
        hashed_password = generate_password_hash(form.password.data, method='sha256')
        new_user = User(username=form.username.data, email=form.email.data, password=hashed_password)
        db.session.add(new_user)
        db.session.commit()
        return render_template("newuser.html", form=form)

    return render_template("signup.html", form=form)


@app.route('/dashboard')
@login_required
def dashboard():
    return render_template("dashboard.html", name=current_user.username)


@app.route('/scheduler', methods=['GET', 'POST'])
@login_required
def scheduler():
    result = Scheduler.query.filter_by(userid=current_user.id).all()
    if request.method == 'POST':
        if int(Scheduler.query.count()) < 6:
            file = request.files['inputfile']
            schedulename = request.form['schedulename']
            data = file.read()
            signature = Scheduler(userid=current_user.id, schedulename=schedulename,schedule=file.filename, data=data)
            db.session.add(signature)
            db.session.commit()
            result = Scheduler.query.filter_by(userid=current_user.id).all()
    return render_template("scheduler.html", result=result)


@app.route('/academics', methods=['POST', 'GET'])
@login_required
def academics():
    grades = Grades.query.filter_by(userid=current_user.id).all()
    result = Mcat.query.filter_by(userid=current_user.id).all()
    result1 = References.query.filter_by(userid=current_user.id).all()
    if request.method == 'POST':
        ogpa = request.form['ogpa']
        sgpa = request.form['sgpa']
        signature = Grades(userid=current_user.id, ogpa=ogpa, sgpa=sgpa)
        db.session.add(signature)
        db.session.commit()
    return render_template("academics.html", result=result, result1=result1, grades=grades)


@app.route('/academicsdetails', methods=['POST', 'GET'])
def academicsdetails():
    result = References.query.filter_by(id=request.form['academicsdetails']).first()
    return render_template("academicsdetails.html", result=result)


@app.route('/academicsdetailsprocess', methods=['POST', 'GET'])
def academicsdetailsprocess():
    edit = References.query.filter_by(id=request.form['update']).first()
    edit.email = request.form['email']
    edit.type = request.form['type']
    edit.status = request.form['status']
    db.session.commit()
    return redirect(url_for('academics'))


@app.route('/mcat', methods=['POST', 'GET'])
def mcat():
    if request.method == 'POST':
        if request.form['examdate'] != '':
            examdate = request.form['examdate']
            overall = request.form['overall']
            cp = request.form['cp']
            cars = request.form['cars']
            bb = request.form['bb']
            ps = request.form['ps']
            signature = Mcat(userid=current_user.id, examdate=examdate, overall=overall, cp=cp, cars=cars, bb=bb, ps=ps)
            db.session.add(signature)
            db.session.commit()
    return redirect(url_for('academics'))


@app.route('/references', methods=['POST', 'GET'])
def references():
    if request.method == 'POST':
        if request.form['name'] != '':
            name = request.form['name']
            email = request.form['email']
            type = request.form['type']
            status = request.form['status']
            signature = References(userid=current_user.id, name=name, email=email, type=type, status=status)
            db.session.add(signature)
            db.session.commit()
    return redirect(url_for('academics'))


@app.route('/activities', methods=['POST', 'GET'])
@login_required
def activities():
    result = Activities.query.filter_by(userid=current_user.id).all()
    if request.method == 'POST':
        if request.form['activity'] != '':
            activity = request.form['activity']
            type = request.form['type']
            reference = request.form['reference']
            hours = request.form['hours']
            signature = Activities(userid=current_user.id, activity=activity, type=type, reference=reference, hours=hours)
            db.session.add(signature)
            db.session.commit()
            result = Activities.query.filter_by(userid=current_user.id).all()
    return render_template("activities.html", result=result)


@app.route('/activitiesdetails', methods=['POST', 'GET'])
def activitiesdetails():
    result = Activities.query.filter_by(id=request.form['activitiesdetails']).first()
    return render_template("activitiesdetails.html", result=result)


@app.route('/activitiesdetailsprocess', methods=['POST', 'GET'])
def activitiesdetailsprocess():
    edit = Activities.query.filter_by(id=request.form['update']).first()
    edit.type = request.form['type']
    edit.hours = request.form['hours']
    edit.reference = request.form['reference']
    edit.startdate = request.form['startdate']
    edit.enddate = request.form['enddate']
    edit.description = request.form['description']
    db.session.commit()
    return redirect(url_for('activities'))


@app.route('/status', methods=['POST', 'GET'])
@login_required
def status():
    result = Status.query.filter_by(userid=current_user.id).all()
    if request.method == 'POST':
        if request.form['university'] != '':
            university = request.form['university']
            primary = request.form['primary']
            secondary = request.form['secondary']
            interview = request.form['interview']
            offer = request.form['offer']
            signature = Status(userid=current_user.id, university=university, primary=primary, secondary=secondary, interview=interview, offer=offer)
            db.session.add(signature)
            db.session.commit()
    return render_template("status.html", result=result)


@app.route('/statusdetails', methods=['POST', 'GET'])
def statusdetails():
    result = Status.query.filter_by(id=request.form['statusdetails']).first()
    return render_template("statusdetails.html", result=result)


@app.route('/statusdetailsprocess', methods=['POST', 'GET'])
def statusdetailsprocess():
    edit = Status.query.filter_by(id=request.form['update']).first()
    edit.primary = request.form['primary']
    edit.secondary = request.form['secondary']
    edit.interview = request.form['interview']
    edit.offer = request.form['offer']
    edit.essay1p = request.form['essay1p']
    edit.essay1a = request.form['essay1a']
    edit.essay2p = request.form['essay2p']
    edit.essay2a = request.form['essay2a']
    edit.essay3p = request.form['essay3p']
    edit.essay3a = request.form['essay3a']
    edit.essay4p = request.form['essay4p']
    edit.essay4a = request.form['essay4a']
    edit.essay5p = request.form['essay5p']
    edit.essay5a = request.form['essay5a']
    db.session.commit()
    return redirect(url_for('status'))


@app.route('/statusdetailsword', methods=['POST', 'GET'])
def statusdetailsword():
    edit = Status.query.filter_by(id=request.form['word']).first()
    document = Document()
    document.add_heading(edit.university, 0)
    #Essay 1
    document.add_heading('Prompt 1:', level=2)
    document.add_paragraph(edit.essay1p)
    document.add_heading('Essay 1:', level=2)
    document.add_paragraph(edit.essay1a)
    #Essay 2
    document.add_heading('Prompt 2:', level=2)
    document.add_paragraph(edit.essay2p)
    document.add_heading('Essay 2:', level=2)
    document.add_paragraph(edit.essay2a)
    #Essay 3
    document.add_heading('Prompt 3:', level=2)
    document.add_paragraph(edit.essay3p)
    document.add_heading('Essay 3:', level=2)
    document.add_paragraph(edit.essay3a)
    #Essay 4
    document.add_heading('Prompt 4:', level=2)
    document.add_paragraph(edit.essay4p)
    document.add_heading('Essay 4:', level=2)
    document.add_paragraph(edit.essay4a)
    #Essay 5
    document.add_heading('Prompt 5:', level=2)
    document.add_paragraph(edit.essay5p)
    document.add_heading('Essay 5:', level=2)
    document.add_paragraph(edit.essay5a)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Essays'
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename='report.doc')


@app.route('/personalstatement', methods=['POST', 'GET'])
@login_required
def personalstatement():
    result = Personal.query.filter_by(userid=current_user.id).all()
    if request.method == 'POST':
        if request.form['title'] != '':
            title = request.form['title']
            signature = Personal(userid=current_user.id, title=title)
            db.session.add(signature)
            db.session.commit()
            result = Personal.query.filter_by(userid=current_user.id).all()
    return render_template("personalstatement.html", result=result)


@app.route('/personalstatementdetails', methods=['POST', 'GET'])
def personalstatementdetails():
    result = Personal.query.filter_by(id=request.form['personalstatementdetails']).first()
    return render_template("personalstatementdetails.html", result=result)


@app.route('/personalstatementdetailsprocess', methods=['POST', 'GET'])
def personalstatementdetailsprocess():
    edit = Personal.query.filter_by(id=request.form['update']).first()
    edit.essay = request.form['essay']
    db.session.commit()
    return redirect(url_for('personalstatement'))


# Deletion Routes

@app.route('/deletescheduler', methods=['POST', 'GET'])
def deletescheduler():
    if request.form['schedulerdelete'] != '':
        Scheduler.query.filter_by(id=int(request.form['schedulerdelete'])).delete()
        db.session.commit()
    return redirect(url_for('scheduler'))


@app.route('/deletegrades', methods=['POST', 'GET'])
def deletegrades():
    if request.form['gradesdelete'] != '':
        Grades.query.filter_by(id=int(request.form['gradesdelete'])).delete()
        db.session.commit()
    return redirect(url_for('academics'))


@app.route('/deletemcat', methods=['POST', 'GET'])
def deletemcat():
    if request.form['mcatdelete'] != '':
        Mcat.query.filter_by(id=int(request.form['mcatdelete'])).delete()
        db.session.commit()
    return redirect(url_for('academics'))


@app.route('/deletereferences', methods=['POST', 'GET'])
def deletereferences():
    if request.form['referencesdelete'] != '':
        References.query.filter_by(id=int(request.form['referencesdelete'])).delete()
        db.session.commit()
    return redirect(url_for('academics'))


@app.route('/deleteactivities', methods=['POST', 'GET'])
def deleteactivities():
    if request.form['activitiesdelete'] != '':
        Activities.query.filter_by(id=int(request.form['activitiesdelete'])).delete()
        db.session.commit()
    return redirect(url_for('activities'))


@app.route('/deletestatus', methods=['POST', 'GET'])
def deletestatus():
    if request.form['statusdelete'] != '':
        Status.query.filter_by(id=int(request.form['statusdelete'])).delete()
        db.session.commit()
    return redirect(url_for('status'))


@app.route('/deletepersonalstatement', methods=['POST', 'GET'])
def deletepersonalstatement():
    if request.form['personalstatementdelete'] != '':
        Personal.query.filter_by(id=int(request.form['personalstatementdelete'])).delete()
        db.session.commit()
    return redirect(url_for('personalstatement'))


# Makes Summary Word Doc
@app.route('/summary')
@login_required
def summary():
    activities = Activities.query.filter_by(userid=current_user.id).all()
    grades = Grades.query.filter_by(userid=current_user.id).all()
    mcat = Mcat.query.filter_by(userid=current_user.id).all()
    references = References.query.filter_by(userid=current_user.id).all()
    status = Status.query.filter_by(userid=current_user.id).all()

    document = Document()
    document.add_heading("Summary", 0)
    document.add_heading('GPA', 1)
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Overall'
    hdr_cells[1].text = 'Science'
    for item in grades:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.ogpa)
        row_cells[1].text = str(item.sgpa)

    document.add_heading('MCAT', 1)
    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Overall'
    hdr_cells[2].text = 'C/P'
    hdr_cells[3].text = 'CARS'
    hdr_cells[4].text = 'B/B'
    hdr_cells[5].text = 'P/S'
    for item in mcat:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.examdate)
        row_cells[1].text = str(item.overall)
        row_cells[2].text = str(item.cp)
        row_cells[3].text = str(item.cars)
        row_cells[4].text = str(item.bb)
        row_cells[5].text = str(item.ps)

    document.add_heading('References', 1)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Email'
    hdr_cells[2].text = 'Type'
    hdr_cells[3].text = 'Status'
    for item in references:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.name)
        row_cells[1].text = str(item.email)
        row_cells[2].text = str(item.type)
        row_cells[3].text = str(item.status)

    document.add_heading('Activities', 1)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Activity'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Reference'
    hdr_cells[3].text = 'Hours'
    for item in activities:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.activity)
        row_cells[1].text = str(item.type)
        row_cells[2].text = str(item.reference)
        row_cells[3].text = str(item.hours)

    document.add_heading('Application Status', 1)
    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'University'
    hdr_cells[1].text = 'Primary'
    hdr_cells[2].text = 'Secondary'
    hdr_cells[3].text = 'Interview'
    hdr_cells[4].text = 'Offer'
    for item in status:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.university)
        row_cells[1].text = str(item.primary)
        row_cells[2].text = str(item.secondary)
        row_cells[3].text = str(item.interview)
        row_cells[4].text = str(item.offer)

    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename='report.doc')


# Flask app initialized
if __name__ == '__main__':
    app.run()